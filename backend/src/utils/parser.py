import os
import logging
import fitz  # PyMuPDF
import docx

logger = logging.getLogger(__name__)

def parse_pdf(file_path: str) -> str:
    """
    Extracts plain text from a PDF document page-by-page.
    """
    logger.info(f"Parsing PDF file: {file_path}")
    text_content = []
    try:
        with fitz.open(file_path) as doc:
            for page_num, page in enumerate(doc):
                page_text = page.get_text()
                if page_text:
                    text_content.append(page_text)
        return "\n\n--- Page Break ---\n\n".join(text_content)
    except Exception as e:
        logger.error(f"Error parsing PDF file {file_path}: {e}")
        raise RuntimeError(f"Failed to parse PDF file: {e}")

def parse_docx(file_path: str) -> str:
    """
    Extracts text from paragraphs and tables in a DOCX document.
    """
    logger.info(f"Parsing DOCX file: {file_path}")
    try:
        doc = docx.Document(file_path)
        text_content = []
        
        # Parse paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
                
        # Parse tables (often contain experience details or skills)
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    text_content.append(" | ".join(row_text))
                    
        return "\n".join(text_content)
    except Exception as e:
        logger.error(f"Error parsing DOCX file {file_path}: {e}")
        raise RuntimeError(f"Failed to parse DOCX file: {e}")

def parse_resume_file(file_path: str) -> str:
    """
    Detects the file extension and routes it to the corresponding parser.
    """
    _, ext = os.path.splitext(file_path.lower())
    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext == ".docx":
        return parse_docx(file_path)
    else:
        logger.error(f"Unsupported file format extension: {ext}")
        raise ValueError(f"Unsupported file format: '{ext}'. Only PDF and DOCX are allowed.")
